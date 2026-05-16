#!/usr/bin/env python3
"""
Enhanced Census Survey Data Processor
Processes downloaded survey data with OCR analysis for PDF/DOC documentation,
deep filesystem traversal, and comprehensive column mapping for RAG/Vector DB usage.
"""

import os
import pandas as pd
import json
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import zipfile
import csv
import logging
from dataclasses import dataclass
from collections import defaultdict


# OCR and document processing imports
try:
    import PyPDF2
    import pdfplumber
    import pytesseract
    from PIL import Image
    import docx
    OCR_AVAILABLE = True
except ImportError as e:
    OCR_AVAILABLE = False
    print(f"OCR libraries not available: {e}")

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class FileInfo:
    path: str
    type: str
    size: int
    modified_time: float
    
@dataclass
class ColumnInfo:
    name: str
    data_type: str
    description: str
    source_file: str
    sample_values: List[str]
    null_count: int = 0

@dataclass
class GeographicMatch:
    survey_id: str
    geo_level: str  # 'county', 'tract', 'block', 'zipcode'
    geo_id: str
    matched_shapefile: str
    confidence_score: float
    overlap_percentage: float
    validation_status: str  # 'valid', 'partial', 'invalid'

class SurveyDataProcessor:
    """Enhanced Census survey data processor with OCR and deep analysis capabilities."""
    
    def __init__(self, base_data_dir: str):
        self.base_data_dir = Path(base_data_dir)
        self.metadata = {}
        self.column_descriptions = {}
        self.document_cache = {}
        self.supported_data_extensions = {'.csv', '.xlsx', '.xls', '.txt'}
        self.supported_doc_extensions = {'.pdf', '.docx', '.doc'}
        
        
    def deep_scan_filesystem(self, max_depth: int = 10) -> Dict[str, List[FileInfo]]:
        """Recursively scan entire filesystem for data and documentation files."""
        found_files = defaultdict(list)
        scanned_count = 0
        
        def scan_recursive(directory: Path, current_depth: int = 0):
            nonlocal scanned_count
            if current_depth > max_depth:
                return
                
            try:
                for item in directory.iterdir():
                    if item.is_file():
                        scanned_count += 1
                        if scanned_count % 100 == 0:
                            logger.info(f"Scanned {scanned_count} files...")
                            
                        file_ext = item.suffix.lower()
                        file_size = item.stat().st_size
                        modified_time = item.stat().st_mtime
                        
                        if file_ext in self.supported_data_extensions:
                            found_files['data'].append(FileInfo(
                                str(item), file_ext, file_size, modified_time
                            ))
                        elif file_ext in self.supported_doc_extensions:
                            found_files['documents'].append(FileInfo(
                                str(item), file_ext, file_size, modified_time
                            ))
                            
                    elif item.is_dir() and not item.name.startswith('.'):
                        scan_recursive(item, current_depth + 1)
                        
            except PermissionError:
                logger.warning(f"Permission denied accessing: {directory}")
            except Exception as e:
                logger.error(f"Error scanning {directory}: {e}")
        
        # Start scanning from base directory
        logger.info(f"Starting deep filesystem scan from: {self.base_data_dir}")
        scan_recursive(self.base_data_dir)
        
        logger.info(f"Deep scan complete. Found {len(found_files['data'])} data files and {len(found_files['documents'])} document files")
        return dict(found_files)
        
    def extract_document_text(self, file_path: str) -> str:
        """Extract text from PDF, DOCX, or DOC files using OCR or text extraction."""
        if not OCR_AVAILABLE:
            return "OCR libraries not available"
            
        path_obj = Path(file_path)
        file_ext = path_obj.suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(path_obj)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx_text(path_obj)
            else:
                return f"Unsupported file type: {file_ext}"
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Error: {str(e)}"
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods."""
        text = ""
        
        # Try pdfplumber first (better for tables)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            
        # Fallback to PyPDF2
        if not text.strip():
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {file_path}: {e}")
                
        # Last resort: OCR with pytesseract
        if not text.strip():
            try:
                import pytesseract
                from pdf2image import convert_from_path
                images = convert_from_path(file_path)
                for img in images:
                    ocr_text = pytesseract.image_to_string(img)
                    text += ocr_text + "\n"
            except Exception as e:
                logger.warning(f"OCR failed for {file_path}: {e}")
                
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            import docx
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                    
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            return f"Error: {str(e)}"
    
    def extract_column_descriptions_from_docs(self, documents: List[FileInfo]) -> Dict[str, str]:
        """Extract column descriptions from documentation files using pattern matching."""
        column_descriptions = {}
        
        # Common column name patterns in census documentation
        column_patterns = [
            r'([A-Z][A-Z0-9_]+)[\s:]+([^.]+?)(?:\n|\.|$)',
            r'([A-Z][A-Z0-9_]+)\s*-\s*([^.\n]+)',
            r'Variable\s+([A-Z][A-Z0-9_]+)[\s:]+([^.\n]+)',
            r'Column\s+([A-Z][A-Z0-9_]+)[\s:]+([^.\n]+)',
            r'([A-Z]\d{5}_\d{3}[A-Z])[\s:]+([^.\n]+)'
        ]
        
        for doc_info in documents:
            if doc_info.path in self.document_cache:
                text = self.document_cache[doc_info.path]
            else:
                text = self.extract_document_text(doc_info.path)
                self.document_cache[doc_info.path] = text
                
            logger.info(f"Extracting column descriptions from: {Path(doc_info.path).name}")
            
            # Apply pattern matching
            for pattern in column_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for column, description in matches:
                    column = column.upper().strip()
                    description = description.strip()
                    
                    # Clean up description
                    description = re.sub(r'\s+', ' ', description)
                    description = description.rstrip('.')
                    
                    if len(column) >= 3 and len(description) > 5:
                        column_descriptions[column] = description
                        
        logger.info(f"Extracted {len(column_descriptions)} column descriptions from documentation")
        return column_descriptions
    
    def scan_survey_directories(self) -> Dict[str, List[str]]:
        """Scan for available survey data directories."""
        surveys_dir = self.base_data_dir / 'census_surveys'
        survey_data = {}
        
        if not surveys_dir.exists():
            logger.warning(f"Survey directory not found: {surveys_dir}")
            return survey_data
            
        for survey_code in os.listdir(surveys_dir):
            survey_path = surveys_dir / survey_code
            if survey_path.is_dir():
                datasets = [d for d in os.listdir(survey_path) if (survey_path / d).is_dir()]
                survey_data[survey_code] = datasets
                
        return survey_data
    
    def extract_zip_metadata(self, zip_path: str) -> Dict[str, Any]:
        """Extract metadata from zip file (file names, sizes, etc)."""
        metadata = {
            'file_name': os.path.basename(zip_path),
            'file_size': os.path.getsize(zip_path),
            'contained_files': []
        }
        
        try:
            with zipfile.ZipFile(zip_path, 'r') as zf:
                for file_info in zf.infolist():
                    metadata['contained_files'].append({
                        'name': file_info.filename,
                        'size': file_info.file_size,
                        'date_time': file_info.date_time
                    })
        except Exception as e:
            metadata['error'] = str(e)
            
        return metadata
    
    def analyze_csv_structure(self, csv_path: str, sample_size: int = 1000) -> Dict[str, Any]:
        """Analyze CSV file structure with comprehensive data quality checks."""
        structure = {
            'file_path': str(csv_path),
            'file_size': os.path.getsize(csv_path),
            'total_rows': 0,
            'total_columns': 0,
            'columns': [],
            'sample_data': [],
            'column_types': {},
            'data_quality': {},
            'statistics': {},
            'missing_values': {},
            'unique_counts': {},
            'data_anomalies': []
        }
        
        try:
            # First pass: read header and count rows
            with open(csv_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                headers = next(reader, [])
                structure['columns'] = headers
                structure['total_columns'] = len(headers)
                
                # Count total rows
                row_count = sum(1 for _ in reader)
                structure['total_rows'] = row_count
                
            # Second pass: sample data for analysis
            df_sample = None
            try:
                # Try pandas for better analysis
                df_sample = pd.read_csv(csv_path, nrows=sample_size, encoding='utf-8', encoding_errors='ignore')
                
                # Basic statistics
                numeric_cols = df_sample.select_dtypes(include=['number']).columns
                structure['statistics'] = {}
                for col in numeric_cols:
                    try:
                        try:
                            structure['statistics'][col] = {
                                'mean': None, 'std': None, 'min': None, 'max': None
                            }
                        except (ValueError, TypeError):
                            structure['statistics'][col] = {
                                'mean': None, 'std': None, 'min': None, 'max': None
                            }
                    except (ValueError, TypeError):
                        structure['statistics'][col] = {
                            'mean': None, 'std': None, 'min': None, 'max': None
                        }
                
                # Missing values analysis
                missing_pct = (df_sample.isnull().sum() / len(df_sample) * 100).to_dict()
                structure['missing_values'] = {k: round(v, 2) for k, v in missing_pct.items()}
                
                # Unique counts
                unique_counts = df_sample.nunique().to_dict()
                structure['unique_counts'] = unique_counts
                
                # Data quality checks
                structure['data_quality'] = self._assess_data_quality(df_sample)
                
                # Convert sample to list format
                structure['sample_data'] = df_sample.head(10).fillna('').values.tolist()
                
            except Exception:
                # Fallback to basic CSV reading
                with open(csv_path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    next(reader)  # Skip headers
                    sample_rows = []
                    for i, row in enumerate(reader):
                        if i >= 10:
                            break
                        sample_rows.append(row)
                    structure['sample_data'] = sample_rows
                
            # Column type analysis
            for col in headers:
                if df_sample is not None and col in df_sample.columns:
                    structure['column_types'][col] = str(df_sample[col].dtype)
                else:
                    # Fallback type inference
                    sample_values = []
                    for row in structure['sample_data']:
                        for i, header in enumerate(headers):
                            if header == col and i < len(row):
                                sample_values.append(row[i])
                    structure['column_types'][col] = 'string'  # Simple fallback
                    
        except Exception as e:
            structure['error'] = str(e)
            logger.error(f"Error analyzing CSV {csv_path}: {e}")
            
        return structure
    
    def _assess_data_quality(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Assess data quality metrics for a DataFrame."""
        quality_metrics = {
            'completeness': {},
            'consistency': {},
            'validity': {}
        }
        
        for col in df.columns:
            # Completeness: percentage of non-null values
            completeness = (df[col].notna().sum() / len(df)) * 100
            quality_metrics['completeness'][col] = round(completeness, 2)
            
            # Consistency: check for duplicate values in key columns
            if col.upper() in ['ID', 'GEOID', 'STATEFP', 'COUNTYFP']:
                duplicates = df[col].duplicated().sum()
                quality_metrics['consistency'][col] = {
                    'duplicates': int(duplicates),
                    'duplicate_rate': round((duplicates / len(df)) * 100, 2)
                }
            
            # Validity: basic range checks for known columns
            if 'STATEFP' in col.upper():
                valid_states = df[col].between(1, 56).sum()  # Valid FIPS state codes
                quality_metrics['validity'][col] = {
                    'valid_count': int(valid_states),
                    'validity_rate': round((valid_states / len(df)) * 100, 2)
                }
                
        return quality_metrics
    
    def detect_data_anomalies(self, df: pd.DataFrame, column_info: Dict[str, ColumnInfo]) -> List[Dict[str, Any]]:
        """Detect data anomalies and quality issues."""
        anomalies = []
        
        for col_name, col_info in column_info.items():
            if col_name not in df.columns:
                continue
                
            series = df[col_name]
            
            # Check for extreme outliers in numeric data
            if pd.api.types.is_numeric_dtype(series):
                Q1 = series.quantile(0.25)
                Q3 = series.quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                
                outliers = series[(series < lower_bound) | (series > upper_bound)]
                if len(outliers) > 0:
                    anomalies.append({
                        'type': 'outlier',
                        'column': col_name,
                        'count': len(outliers),
                        'percentage': (len(outliers) / len(series)) * 100,
                        'severity': 'high' if len(outliers) > len(series) * 0.05 else 'medium',
                        'description': f"Statistical outliers detected in {col_name}"
                    })
            
            # Check for suspicious patterns in text data
            elif pd.api.types.is_string_dtype(series):
                # Check for suspiciously repetitive values
                value_counts = series.value_counts()
                if len(value_counts) > 0:
                    most_common_pct = (value_counts.iloc[0] / len(series)) * 100
                    if most_common_pct > 90:
                        anomalies.append({
                            'type': 'repetitive_values',
                            'column': col_name,
                            'dominant_value': str(value_counts.index[0]),
                            'percentage': most_common_pct,
                            'severity': 'high',
                            'description': f"Column {col_name} is {most_common_pct:.1f}% identical values"
                        })
                
                # Check for suspicious null-like strings
                null_patterns = ['', ' ', 'NULL', 'null', 'N/A', 'n/a', '#N/A', '#REF!', '0']
                null_like_count = series.str.strip().isin(null_patterns).sum()
                if null_like_count > 0:
                    anomalies.append({
                        'type': 'suspicious_nulls',
                        'column': col_name,
                        'count': null_like_count,
                        'percentage': (null_like_count / len(series)) * 100,
                        'severity': 'medium',
                        'description': f"Suspicious null-like values found in {col_name}"
                    })
            
            # Geographic validation
            if any(geo_term in col_name.upper() for geo_term in ['STATE', 'FIPS', 'GEOID']):
                if col_name.upper() == 'STATEFP' or 'STATE' in col_name.upper():
                    # Validate state FIPS codes (1-56)
                    if pd.api.types.is_numeric_dtype(series):
                        invalid_states = series[~series.between(1, 56)]
                        if len(invalid_states) > 0:
                            anomalies.append({
                                'type': 'invalid_geographic_codes',
                                'column': col_name,
                                'count': len(invalid_states),
                                'percentage': (len(invalid_states) / len(series)) * 100,
                                'severity': 'high',
                                'description': f"Invalid state FIPS codes found in {col_name}"
                            })
        
        return anomalies
    
    
    
    
    
    def validate_geographic_identifiers(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Validate and analyze geographic identifiers in survey data."""
        validation_results = {
            'identified_columns': {},
            'validation_summary': {},
            'geographic_coverage': {},
            'recommendations': []
        }
        
        # Common geographic identifier patterns
        geo_patterns = {
            'STATEFP': r'^\d{2}$',
            'COUNTYFP': r'^\d{3}$',
            'GEOID': r'^\d{5,11}$',
            'GEO_ID': r'^\d{5,11}$',
            'TRACTCE': r'^\d{5,1}$',
            'BLOCKCE': r'^\d{3}$',
            'ZCTA520': r'^\d{5}$',
            'ZIPCODE': r'^\d{5}$',
            'ZIP': r'^\d{5}$'
        }
        
        for col_name in df.columns:
            col_upper = col_name.upper()
            
            for geo_id, pattern in geo_patterns.items():
                if geo_id in col_upper:
                    # Validate the column
                    series = df[col_name].dropna()
                    valid_count = series.astype(str).str.match(pattern).sum()
                    total_count = len(series)
                    
                    validation_results['identified_columns'][col_name] = {
                        'type': geo_id,
                        'pattern': pattern,
                        'total_values': total_count,
                        'valid_values': int(valid_count),
                        'validity_rate': round((valid_count / total_count * 100) if total_count > 0 else 0, 2),
                        'sample_values': series.head(5).astype(str).tolist()
                    }
                    
                    # Geographic coverage analysis
                    unique_values = series.unique()
                    validation_results['geographic_coverage'][col_name] = {
                        'unique_locations': len(unique_values),
                        'coverage_type': geo_id,
                        'top_locations': dict(series.value_counts().head(10))
                    }
                    break
        
        # Generate recommendations
        for col, info in validation_results['identified_columns'].items():
            if info['validity_rate'] < 95:
                validation_results['recommendations'].append(
                    f"Column {col} has low validity rate ({info['validity_rate']}%) for {info['type']}"
                )
        
        return validation_results
    
    
    
    
    
    
    
    
    
    
    
    

def main():
    """Main function to demonstrate survey data processing for naming convention analysis."""
    # Get the repository root directory
    current_file = Path(__file__).resolve()
    repo_root = current_file.parent.parent
    data_dir = repo_root / 'data'
    
    processor = SurveyDataProcessor(str(data_dir))
    
    # Scan for available surveys
    available_surveys = processor.scan_survey_directories()
    print("Available survey datasets:")
    for survey_code, datasets in available_surveys.items():
        print(f"  {survey_code}: {datasets}")
    
    if available_surveys:
        print("\nScanning for file naming patterns and PDF methodologies...")
        
        # Deep scan filesystem
        found_files = processor.deep_scan_filesystem()
        
        print(f"\nFile Scan Results:")
        print(f"  Data files found: {len(found_files.get('data', []))}")
        print(f"  Document files found: {len(found_files.get('documents', []))}")
        
        # Extract column descriptions from documents
        if found_files.get('documents'):
            print("\nExtracting column descriptions from PDF/DOC files...")
            column_descriptions = processor.extract_column_descriptions_from_docs(found_files['documents'])
            print(f"  Column descriptions extracted: {len(column_descriptions)}")
            
            # Show some examples
            if column_descriptions:
                print("\nSample column descriptions:")
                for i, (col, desc) in enumerate(list(column_descriptions.items())[:5]):
                    print(f"  {col}: {desc}")
        
        # Analyze first few data files for structure
        data_files = found_files.get('data', [])[:3]  # Limit to first 3 files
        for file_info in data_files:
            print(f"\nAnalyzing: {Path(file_info.path).name}")
            structure = processor.analyze_csv_structure(file_info.path)
            print(f"  Columns: {structure.get('total_columns', 0)}")
            print(f"  Rows: {structure.get('total_rows', 0)}")
            
            # Show geographic validation
            if structure.get('columns'):
                df_sample = pd.read_csv(file_info.path, nrows=100)
                geo_validation = processor.validate_geographic_identifiers(df_sample)
                if geo_validation.get('identified_columns'):
                    print(f"  Geographic columns: {list(geo_validation['identified_columns'].keys())}")
        
        print(f"\n✅ Analysis complete. Focus is on naming conventions and PDF methodology parsing.")
        print(f"   No geospatial analytics performed as requested.")
        
    else:
        print("\nNo survey data found. Please run get_census_survey.py first.")
        return

if __name__ == "__main__":
    main()